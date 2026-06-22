#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
LifeOS Thesis — Professional Word Document Generator
Reads thesis.md and produces a fully styled thesis.docx
Design inspired by: docs/presentation/styles.css
  - Accent colour  : #4F46E5 (Indigo)
  - Accent light   : #6366F1
  - Success        : #10B981 (Emerald)
  - Warning        : #F59E0B (Amber)
  - Font family    : Calibri
"""

import os
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx.enum.text as docx_breaks

# ─── Colour Palette (from presentation/styles.css) ────────────────────────────
ACCENT      = RGBColor(0x4F, 0x46, 0xE5)   # #4F46E5  Indigo
ACCENT_LIGHT= RGBColor(0x63, 0x66, 0xF1)   # #6366F1
SUCCESS     = RGBColor(0x10, 0xB9, 0x81)   # #10B981  Emerald
WARNING     = RGBColor(0xF5, 0x9E, 0x0B)   # #F59E0B  Amber
DARK_TEXT   = RGBColor(0x11, 0x18, 0x27)   # #111827
MID_GRAY    = RGBColor(0x6B, 0x72, 0x80)   # #6B7280
LIGHT_GRAY  = RGBColor(0xF3, 0xF4, 0xF6)   # #F3F4F6
CODE_BG     = RGBColor(0xF8, 0xF9, 0xFA)   # near-white code block bg
TABLE_HDR   = RGBColor(0x4F, 0x46, 0xE5)   # table header fill same as accent
TABLE_HDR_TXT=RGBColor(0xFF, 0xFF, 0xFF)   # white text on header
TABLE_ALT   = RGBColor(0xF0, 0xEF, 0xFF)   # light lavender alternating row
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
BLACK       = RGBColor(0x00, 0x00, 0x00)
BORDER_CLR  = RGBColor(0xE5, 0xE7, 0xEB)   # #E5E7EB

FONT_BODY   = "Calibri"
FONT_HEAD   = "Calibri"
FONT_CODE   = "Courier New"

# ─── Helper: XML utilities ────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Set table cell background colour via raw XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set inner cell padding in dxa (twips)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), val.get("val", "single"))
            el.set(qn("w:sz"), val.get("sz", "4"))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), val.get("color", "000000"))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def para_border_bottom(para, color="4F46E5", sz="12"):
    """Add a bottom border (rule line) to a paragraph."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_para_shading(para, fill_hex: str):
    """Set paragraph background shading."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)

def set_run_highlight(run, color_hex: str):
    """Set run background (character shading)."""
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    rPr.append(shd)

def add_page_break(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(docx_breaks.WD_BREAK.PAGE)

def set_page_margins(section, top=2.54, bottom=2.54, left=3.17, right=2.54):
    section.top_margin    = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin   = Cm(left)
    section.right_margin  = Cm(right)

def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

# ─── Style Builders ──────────────────────────────────────────────────────────

def configure_styles(doc: Document):
    styles = doc.styles

    # Normal / Body text
    normal = styles["Normal"]
    normal.font.name  = FONT_BODY
    normal.font.size  = Pt(11)
    normal.font.color.rgb = DARK_TEXT
    normal.paragraph_format.space_after  = Pt(8)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.5

    # Heading 1 — Chapter titles
    h1 = styles["Heading 1"]
    h1.font.name  = FONT_HEAD
    h1.font.size  = Pt(20)
    h1.font.bold  = True
    h1.font.color.rgb = ACCENT
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after  = Pt(12)
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.page_break_before = True

    # Heading 2 — Section titles
    h2 = styles["Heading 2"]
    h2.font.name  = FONT_HEAD
    h2.font.size  = Pt(15)
    h2.font.bold  = True
    h2.font.color.rgb = DARK_TEXT
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after  = Pt(8)
    h2.paragraph_format.keep_with_next = True

    # Heading 3 — Sub-section titles
    h3 = styles["Heading 3"]
    h3.font.name  = FONT_HEAD
    h3.font.size  = Pt(12.5)
    h3.font.bold  = True
    h3.font.color.rgb = ACCENT
    h3.paragraph_format.space_before = Pt(14)
    h3.paragraph_format.space_after  = Pt(6)
    h3.paragraph_format.keep_with_next = True

    # Heading 4
    h4 = styles["Heading 4"]
    h4.font.name  = FONT_HEAD
    h4.font.size  = Pt(11)
    h4.font.bold  = True
    h4.font.italic = True
    h4.font.color.rgb = DARK_TEXT
    h4.paragraph_format.space_before = Pt(10)
    h4.paragraph_format.space_after  = Pt(4)

def styled_heading(doc, text: str, level: int):
    style_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3", 4: "Heading 4"}
    style = style_map.get(level, "Heading 2")
    p = doc.add_heading(text, level=level)
    p.style = doc.styles[style]
    if level == 1:
        para_border_bottom(p, color="4F46E5", sz="8")
    return p

# ─── Page Sections ───────────────────────────────────────────────────────────

def add_title_page(doc: Document):
    section = doc.sections[0]
    set_page_margins(section)

    def center_para(text, size, bold=False, color=DARK_TEXT, space_before=0, space_after=6, italic=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        run.font.name  = FONT_HEAD
        run.font.size  = Pt(size)
        run.font.bold  = bold
        run.font.italic = italic
        run.font.color.rgb = color
        return p

    # Top rule
    top_rule = doc.add_paragraph()
    top_rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    top_rule.paragraph_format.space_before = Pt(0)
    top_rule.paragraph_format.space_after  = Pt(0)
    para_border_bottom(top_rule, color="4F46E5", sz="24")

    center_para("", 6, space_before=30)  # spacer
    center_para("DEPARTMENT OF COMPUTER APPLICATION", 10, bold=True, color=MID_GRAY, space_after=2)
    center_para("Patna College, Patna", 11, bold=False, color=DARK_TEXT, space_after=30)

    # Main title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(10)
    p_title.paragraph_format.space_after  = Pt(4)
    p_title.paragraph_format.line_spacing = 1.2
    r1 = p_title.add_run("LifeOS")
    r1.font.name = FONT_HEAD; r1.font.size = Pt(42); r1.font.bold = True
    r1.font.color.rgb = ACCENT
    r2 = p_title.add_run(": Smart Life Decision Operating System")
    r2.font.name = FONT_HEAD; r2.font.size = Pt(28); r2.font.bold = True
    r2.font.color.rgb = DARK_TEXT

    center_para("A Project Report Submitted in Partial Fulfilment of the\nRequirements for the Degree of", 11, color=MID_GRAY, space_before=8, space_after=4)
    center_para("Bachelor of Computer Applications (BCA)", 14, bold=True, color=DARK_TEXT, space_after=20)

    # Horizontal rule
    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_border_bottom(rule, color="4F46E5", sz="8")
    rule.paragraph_format.space_after = Pt(16)

    center_para("Submitted by", 10, color=MID_GRAY, bold=True, space_before=0, space_after=2)
    center_para("Badal Kumar", 26, bold=True, color=DARK_TEXT, space_before=0, space_after=4)

    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.space_after = Pt(20)
    r = p_meta.add_run("Roll No: 10   |   BCA 6th Semester   |   Session: 2023–2026")
    r.font.name = FONT_BODY; r.font.size = Pt(11); r.font.color.rgb = MID_GRAY

    # Guide and HOD
    rule2 = doc.add_paragraph()
    rule2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_border_bottom(rule2, color="E5E7EB", sz="6")
    rule2.paragraph_format.space_after = Pt(14)

    p_guide = doc.add_paragraph()
    p_guide.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_guide.paragraph_format.space_after = Pt(4)
    rg1 = p_guide.add_run("Project Guide: "); rg1.font.name = FONT_BODY; rg1.font.size = Pt(11); rg1.font.color.rgb = MID_GRAY
    rg2 = p_guide.add_run("Ajit Sir"); rg2.font.name = FONT_BODY; rg2.font.size = Pt(11); rg2.font.bold = True; rg2.font.color.rgb = DARK_TEXT

    p_hod = doc.add_paragraph()
    p_hod.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_hod.paragraph_format.space_after = Pt(24)
    rh1 = p_hod.add_run("Coordinator & HOD: "); rh1.font.name = FONT_BODY; rh1.font.size = Pt(11); rh1.font.color.rgb = MID_GRAY
    rh2 = p_hod.add_run("Dr. Dinesh Manjhi"); rh2.font.name = FONT_BODY; rh2.font.size = Pt(11); rh2.font.bold = True; rh2.font.color.rgb = DARK_TEXT

    center_para("June 2026", 12, bold=True, color=ACCENT, space_before=10)

    # Bottom rule
    bot_rule = doc.add_paragraph()
    bot_rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bot_rule.paragraph_format.space_before = Pt(10)
    para_border_bottom(bot_rule, color="4F46E5", sz="24")


def add_certificate(doc):
    doc.add_page_break()
    p = doc.add_heading("CERTIFICATE", 1)
    p.paragraph_format.page_break_before = False
    para_border_bottom(p, color="4F46E5", sz="6")
    doc.add_paragraph()

    body = (
        "This is to certify that the project entitled "
        "\"LifeOS: Smart Life Decision Operating System\" "
        "submitted by Badal Kumar (Roll No. 10), BCA 6th Semester, "
        "Session 2023–2026, Department of Computer Application, "
        "Patna College, Patna, is a record of bonafide work carried out "
        "by the candidate under my supervision and guidance.\n\n"
        "The project report has not been submitted previously for the award "
        "of any other degree or diploma of this or any other university or institution."
    )
    p = doc.add_paragraph(body, style="Normal")

    doc.add_paragraph()
    doc.add_paragraph()

    # Signature line table
    t = doc.add_table(1, 2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    left_cell = t.rows[0].cells[0]
    right_cell = t.rows[0].cells[1]
    for cell, title, name in [
        (left_cell, "Project Guide", "Ajit Sir"),
        (right_cell, "Coordinator & HOD", "Dr. Dinesh Manjhi"),
    ]:
        cell.paragraphs[0].text = ""
        p = cell.add_paragraph()
        r = p.add_run(f"\n\n\n___________________\n{title}\n")
        r.font.name = FONT_BODY; r.font.size = Pt(11)
        r = p.add_run(name)
        r.font.name = FONT_BODY; r.font.size = Pt(11); r.font.bold = True
        r = p.add_run("\nDept. of Computer Application\nPatna College, Patna")
        r.font.name = FONT_BODY; r.font.size = Pt(11)

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("Date: _________________________       Place: Patna")
    r.font.name = FONT_BODY; r.font.size = Pt(11)


def add_declaration(doc):
    doc.add_page_break()
    p = doc.add_heading("DECLARATION", 1)
    p.paragraph_format.page_break_before = False
    para_border_bottom(p, color="4F46E5", sz="6")
    doc.add_paragraph()

    body = (
        "I, Badal Kumar, student of Bachelor of Computer Applications (BCA), "
        "6th Semester, Roll No. 10, Session 2023–2026, Patna College, Patna, "
        "hereby declare that the project report entitled \"LifeOS: Smart Life Decision "
        "Operating System\" submitted to the Department of Computer Application is my "
        "own original work.\n\n"
        "This project has not been submitted, either in part or in full, to any other "
        "institution or university for the award of any other degree or diploma. All sources "
        "referred to in the preparation of this report have been duly acknowledged."
    )
    doc.add_paragraph(body, style="Normal")
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    r = p.add_run("\n\n\n___________________\nSignature of Candidate\n")
    r.font.name = FONT_BODY; r.font.size = Pt(11)
    r = p.add_run("Name: Badal Kumar\n")
    r.font.name = FONT_BODY; r.font.size = Pt(11); r.font.bold = True
    r = p.add_run("Roll No: 10\nDate: _________________________\nPlace: Patna")
    r.font.name = FONT_BODY; r.font.size = Pt(11)


def add_acknowledgement(doc):
    doc.add_page_break()
    p = doc.add_heading("ACKNOWLEDGEMENT", 1)
    p.paragraph_format.page_break_before = False
    para_border_bottom(p, color="4F46E5", sz="6")
    doc.add_paragraph()

    paras = [
        "Completing this project has been one of the most rewarding experiences of my BCA journey, and it would not have been possible without the support of many people.",
        "I want to begin by expressing my deepest gratitude to Dr. Dinesh Manjhi, Head of the Department of Computer Application, Patna College, Patna. His encouragement and administrative guidance made the environment for this project possible.",
        "I am sincerely grateful to my project guide, Ajit Sir, for his time, patience, and direction throughout the development process. His technical feedback helped me avoid many mistakes and pushed me to think about problems more carefully. I genuinely could not have completed a project of this scope without that guidance.",
        "To all the faculty members of the Department of Computer Application — thank you for three years of teaching that gave me the foundation to take on a project like this.",
        "To my family — thank you for your support and for tolerating the long nights and weekends spent in front of a screen. Your belief in me made the difficult parts easier.",
        "And to everyone who took the time to test the application and share feedback — your honest opinions made LifeOS better than it would have been otherwise.",
    ]
    for para_text in paras:
        doc.add_paragraph(para_text, style="Normal")

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("Badal Kumar\nBCA 6th Semester, Roll No. 10\nPatna College, Patna")
    r.font.name = FONT_BODY; r.font.size = Pt(11); r.font.bold = True


def add_abstract(doc):
    doc.add_page_break()
    p = doc.add_heading("ABSTRACT", 1)
    p.paragraph_format.page_break_before = False
    para_border_bottom(p, color="4F46E5", sz="6")
    doc.add_paragraph()

    abstract_paras = [
        "Every day, people make thousands of decisions — from the small and routine to the large and life-changing. Yet almost no one goes back to study those decisions afterward: whether they worked, what went wrong, and what could be learned. Over time, this leads to repeating the same mistakes and never truly improving as a decision-maker.",
        "LifeOS — Smart Life Decision Operating System is a mobile application designed to solve this problem. It gives users a structured way to log important decisions, track how those decisions turn out over time, and get advice from an AI that actually knows their personal decision history.",
        "The application is a full-stack mobile system built with React Native (Expo SDK 55) for the mobile frontend, Node.js with Express v5 (running on the Bun runtime) for the backend, and PostgreSQL with the pgvector extension as the database. AI functionality is powered by the Groq API using a custom 4-layer context assembly pipeline that personalises every response to the individual user.",
        "The core features are: a structured decision journal, scheduled outcome check-ins with automated reminders, a conversational AI advisor with real-time Server-Sent Events (SSE) streaming, automatic behavioural pattern detection that runs in the background, and a personal memory system where the AI stores and recalls important facts about the user across sessions.",
        "In terms of scale, the project includes over 20 database tables, 40+ REST API endpoints, 10+ mobile screens, 25+ React Native components, and approximately 11,000 lines of TypeScript code. The architecture demonstrates real-world engineering practices including JWT authentication with token rotation, type-safe database access via Drizzle ORM, semantic vector search using HNSW indexing, asynchronous AI processing, and encrypted mobile token storage.",
    ]
    for para_text in abstract_paras:
        doc.add_paragraph(para_text, style="Normal")

    doc.add_paragraph()
    kw_para = doc.add_paragraph()
    r1 = kw_para.add_run("Keywords: ")
    r1.font.bold = True; r1.font.name = FONT_BODY; r1.font.size = Pt(11)
    r2 = kw_para.add_run(
        "React Native, Expo, Node.js, Express, PostgreSQL, pgvector, Drizzle ORM, "
        "JWT Authentication, Retrieval-Augmented Generation (RAG), Groq API, "
        "Server-Sent Events, Decision Support System, BCA Final Year Project"
    )
    r2.font.name = FONT_BODY; r2.font.size = Pt(11); r2.font.italic = True


def add_toc(doc):
    doc.add_page_break()
    p = doc.add_heading("TABLE OF CONTENTS", 1)
    p.paragraph_format.page_break_before = False
    para_border_bottom(p, color="4F46E5", sz="6")
    doc.add_paragraph()

    toc_items = [
        ("Certificate", "ii", 0), ("Declaration", "iii", 0), ("Acknowledgement", "iv", 0),
        ("Abstract", "v", 0), ("List of Figures", "viii", 0), ("List of Tables", "ix", 0),
        ("CHAPTER 1: Introduction", "1", 0),
        ("1.1  Background and Motivation", "2", 1), ("1.2  Problem Statement", "3", 1),
        ("1.3  Objectives of the Project", "4", 1), ("1.4  Scope of the Project", "5", 1),
        ("1.5  Report Organisation", "6", 1),
        ("CHAPTER 2: Literature Review", "7", 0),
        ("2.1  Existing Decision Support Systems", "8", 1),
        ("2.2  Behavioural Science and Decision Making", "9", 1),
        ("2.3  AI-Powered Personal Assistants", "11", 1),
        ("2.4  Mobile Self-Improvement Apps", "12", 1), ("2.5  Research Gap", "13", 1),
        ("CHAPTER 3: System Analysis", "14", 0),
        ("3.1  Feasibility Study", "15", 1), ("3.2  Requirement Analysis", "18", 1),
        ("3.3  Use Case Analysis", "22", 1), ("3.4  Data Flow Diagram", "28", 1),
        ("CHAPTER 4: System Design", "32", 0),
        ("4.1  Architecture Overview", "33", 1),
        ("4.2  Entity Relationship Diagram", "39", 1),
        ("4.3  Database Design", "45", 1), ("4.4  API Design", "58", 1),
        ("4.5  UI/UX Design", "67", 1), ("4.6  AI Intelligence Architecture", "74", 1),
        ("CHAPTER 5: Implementation", "80", 0),
        ("5.1  Technology Stack", "81", 1), ("5.2  Backend Implementation", "85", 1),
        ("5.3  Frontend Implementation", "100", 1), ("5.4  AI Integration", "115", 1),
        ("5.5  Security Implementation", "122", 1),
        ("CHAPTER 6: Testing", "128", 0),
        ("6.1  Testing Strategy", "129", 1), ("6.2  API Testing", "131", 1),
        ("6.3  Functional Testing", "135", 1), ("6.4  Test Results Summary", "138", 1),
        ("CHAPTER 7: Results and Discussion", "140", 0),
        ("7.1  Project Metrics", "141", 1), ("7.2  Objectives Achieved", "142", 1),
        ("7.3  Challenges and Solutions", "145", 1),
        ("CHAPTER 8: Future Scope", "150", 0),
        ("CHAPTER 9: Conclusion", "157", 0),
        ("References", "161", 0), ("Appendix A: Selected Source Code", "164", 0),
        ("Appendix B: API Reference", "174", 0),
    ]

    for item, page, level in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        indent = Cm(level * 0.8)
        p.paragraph_format.left_indent  = indent

        is_chapter = "CHAPTER" in item
        r = p.add_run(item)
        r.font.name  = FONT_BODY
        r.font.size  = Pt(11 if is_chapter else 10.5)
        r.font.bold  = is_chapter
        r.font.color.rgb = ACCENT if is_chapter else DARK_TEXT

        # Add tab + page number
        tab_run = p.add_run("\t" + page)
        tab_run.font.name  = FONT_BODY
        tab_run.font.size  = Pt(10.5)
        tab_run.font.color.rgb = MID_GRAY

        # Right-align the tab stop for page numbers
        pPr = p._p.get_or_add_pPr()
        tabs_el = OxmlElement("w:tabs")
        tab_el  = OxmlElement("w:tab")
        tab_el.set(qn("w:val"), "right")
        tab_el.set(qn("w:leader"), "dot")
        tab_el.set(qn("w:pos"), "8640")
        tabs_el.append(tab_el)
        pPr.append(tabs_el)

def add_list_of_figures_page(doc):
    doc.add_page_break()
    p = doc.add_heading("LIST OF FIGURES", 1)
    p.paragraph_format.page_break_before = False
    para_border_bottom(p, color="4F46E5", sz="6")
    doc.add_paragraph()
    
    figures = [
        ("Figure 3.1", "Context Diagram — Level 0 DFD", "29"),
        ("Figure 3.2", "Level 1 Data Flow Diagram", "31"),
        ("Figure 4.1", "High-Level System Architecture", "34"),
        ("Figure 4.2", "Detailed Layered Architecture", "37"),
        ("Figure 4.3", "Entity Relationship Diagram — Core Tables", "41"),
        ("Figure 4.4", "Entity Relationship Diagram — AI and Analytics", "43"),
        ("Figure 4.5", "AI Pipeline — 4-Layer Context Assembly", "76"),
        ("Figure 4.6", "Sequence Diagram — Decision Creation Flow", "78"),
        ("Figure 5.1", "Mobile App Folder Structure", "83"),
        ("Figure 5.2", "Backend Folder Structure", "84"),
        ("Figure 5.3", "Login Screen", "[Screenshot]"),
        ("Figure 5.4", "Register Screen", "[Screenshot]"),
        ("Figure 5.5", "Dashboard Screen", "[Screenshot]"),
        ("Figure 5.6", "Decision List Screen", "[Screenshot]"),
        ("Figure 5.7", "Decision Wizard — Step 1", "[Screenshot]"),
        ("Figure 5.8", "Decision Wizard — Step 2", "[Screenshot]"),
        ("Figure 5.9", "Decision Wizard — Step 3", "[Screenshot]"),
        ("Figure 5.10", "Decision Detail Screen", "[Screenshot]"),
        ("Figure 5.11", "Outcome Check-In Screen", "[Screenshot]"),
        ("Figure 5.12", "AI Advisor Chat Screen", "[Screenshot]"),
        ("Figure 5.13", "AI Chat History Screen", "[Screenshot]"),
    ]
    
    for num, title, page in figures:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        
        r1 = p.add_run(f"{num}: ")
        r1.font.bold = True
        r1.font.color.rgb = ACCENT
        
        r2 = p.add_run(title)
        r2.font.color.rgb = DARK_TEXT
        
        tab_run = p.add_run("\t" + page)
        tab_run.font.color.rgb = MID_GRAY
        
        pPr = p._p.get_or_add_pPr()
        tabs_el = OxmlElement("w:tabs")
        tab_el  = OxmlElement("w:tab")
        tab_el.set(qn("w:val"), "right")
        tab_el.set(qn("w:leader"), "dot")
        tab_el.set(qn("w:pos"), "8640")
        tabs_el.append(tab_el)
        pPr.append(tabs_el)

def add_list_of_tables_page(doc):
    doc.add_page_break()
    p = doc.add_heading("LIST OF TABLES", 1)
    p.paragraph_format.page_break_before = False
    para_border_bottom(p, color="4F46E5", sz="6")
    doc.add_paragraph()
    
    tables = [
        ("Table 3.1", "Technical Feasibility Assessment", "16"),
        ("Table 3.2", "Development Timeline", "18"),
        ("Table 3.3", "Functional Requirements", "19"),
        ("Table 3.4", "Non-Functional Requirements", "21"),
        ("Table 4.1", "Database Modules Overview", "46"),
        ("Table 4.2", "users Table — Column Definitions", "47"),
        ("Table 4.3", "user_profiles Table — Column Definitions", "49"),
        ("Table 4.4", "decisions Table — Column Definitions", "50"),
        ("Table 4.5", "outcomes Table — Column Definitions", "53"),
        ("Table 4.6", "outcome_reminders Table — Column Definitions", "55"),
        ("Table 4.7", "ai_chat_sessions Table — Column Definitions", "56"),
        ("Table 4.8", "user_memories Table — Column Definitions", "57"),
        ("Table 4.9", "Authentication API Endpoints", "59"),
        ("Table 4.10", "Decision Management API Endpoints", "61"),
        ("Table 4.11", "Outcome API Endpoints", "62"),
        ("Table 4.12", "Analytics API Endpoints", "63"),
        ("Table 4.13", "AI Module API Endpoints", "64"),
        ("Table 4.14", "Color Palette", "68"),
        ("Table 5.1", "Technology Stack Summary", "81"),
        ("Table 5.2", "Screen Implementation Status", "100"),
        ("Table 6.1", "Authentication Test Cases", "132"),
        ("Table 6.2", "Decision Management Test Cases", "133"),
        ("Table 6.3", "Outcome and AI Test Cases", "134"),
        ("Table 6.4", "Functional Test Results", "138"),
        ("Table 6.5", "Overall Test Summary", "139"),
        ("Table 7.1", "Project Metrics Summary", "141"),
        ("Table 7.2", "Objectives vs. Achievement", "143"),
        ("Table 7.3", "Challenges and Solutions", "146"),
    ]
    
    for num, title, page in tables:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        
        r1 = p.add_run(f"{num}: ")
        r1.font.bold = True
        r1.font.color.rgb = ACCENT
        
        r2 = p.add_run(title)
        r2.font.color.rgb = DARK_TEXT
        
        tab_run = p.add_run("\t" + page)
        tab_run.font.color.rgb = MID_GRAY
        
        pPr = p._p.get_or_add_pPr()
        tabs_el = OxmlElement("w:tabs")
        tab_el  = OxmlElement("w:tab")
        tab_el.set(qn("w:val"), "right")
        tab_el.set(qn("w:leader"), "dot")
        tab_el.set(qn("w:pos"), "8640")
        tabs_el.append(tab_el)
        pPr.append(tabs_el)

# ─── Content Adders ──────────────────────────────────────────────────────────

def add_body_para_parser(doc, text):
    p = doc.add_paragraph(style="Normal")
    add_formatted_text(p, text)
    return p

def add_bullet_parser(doc, text, level=0, checkmark=False):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.0 + level * 0.6)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if checkmark:
        run = p.add_run("✓ ")
        run.font.name = FONT_BODY
        run.font.size = Pt(10.5)
        run.font.color.rgb = SUCCESS
        run.font.bold = True
    add_formatted_text(p, text, font_size=Pt(10.5))
    return p

def add_numbered_list_parser(doc, text, num):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run_num = p.add_run(f"{num}. ")
    run_num.font.name = FONT_BODY
    run_num.font.bold = True
    run_num.font.size = Pt(10.5)
    run_num.font.color.rgb = ACCENT
    add_formatted_text(p, text, font_size=Pt(10.5))
    return p

def add_code_block_parser(doc, lines):
    # Set spacing tight to look like a single code block
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1.5) if i > 0 else Pt(8)
        p.paragraph_format.space_after = Pt(1.5) if i < len(lines)-1 else Pt(8)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.line_spacing = 1.0
        set_para_shading(p, "F8F9FA")
        
        if line == "":
            run = p.add_run(" ")
        else:
            run = p.add_run(line)
        run.font.name = FONT_CODE
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

def add_note_box(doc, text: str, label="NOTE"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    set_para_shading(p, "EEF2FF")
    
    label_run = p.add_run(f"{label}: ")
    label_run.font.name = FONT_BODY
    label_run.font.size = Pt(10)
    label_run.font.bold = True
    label_run.font.color.rgb = ACCENT
    
    add_formatted_text(p, text, font_size=Pt(10), default_color=DARK_TEXT)

def add_blockquote_parser(doc, lines):
    cleaned_lines = []
    for line in lines:
        stripped = line.strip()
        if stripped.startswith('>'):
            stripped = stripped[1:].strip()
        cleaned_lines.append(stripped)
    text = " ".join(cleaned_lines)
    
    label = "NOTE"
    body = text
    match = re.match(r'^\*\*\[?([A-Za-z0-9\s\!]+)\]?\*\*\s*:\s*(.*)', text)
    if match:
        label = match.group(1)
        body = match.group(2)
    elif text.startswith('[!NOTE]') or text.startswith('[!IMPORTANT]') or text.startswith('[!TIP]') or text.startswith('[!WARNING]') or text.startswith('[!CAUTION]'):
        idx = text.find(']')
        label = text[2:idx]
        body = text[idx+1:].strip()
        
    add_note_box(doc, body, label=label.upper())

def add_professional_table(doc, headers: list, rows: list, col_widths=None, alternating=True):
    n_cols = len(headers)
    table = doc.add_table(rows=1, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(hdr_cells[i], "4F46E5")
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=150, right=150)
        for para in hdr_cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.line_spacing = 1.0
            para.paragraph_format.space_after = Pt(0)
            for run in para.runs:
                run.font.name  = FONT_HEAD
                run.font.size  = Pt(10)
                run.font.bold  = True
                run.font.color.rgb = WHITE

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.add_row()
        row.height = Twips(350)
        fill = "F0EFFF" if (alternating and r_idx % 2 == 0) else "FFFFFF"
        for c_idx, cell_text in enumerate(row_data):
            if c_idx >= n_cols:
                break
            cell = row.cells[c_idx]
            cell.text = ""
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_bg(cell, fill)
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(0)
            add_formatted_text(p, cell_text, font_size=Pt(9.5))

    # Column widths
    if col_widths:
        for row in table.rows:
            for j, cell in enumerate(row.cells):
                if j < len(col_widths):
                    cell.width = Cm(col_widths[j])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    return table

def add_table_parser(doc, lines):
    rows_data = []
    for line in lines:
        stripped = line.strip()
        if not stripped.startswith('|'):
            continue
        cells = [c.strip() for c in stripped.split('|')[1:-1]]
        if all(re.match(r'^:?-+:?$', c) for c in cells):
            continue
        rows_data.append(cells)
        
    if not rows_data:
        return
        
    headers = rows_data[0]
    rows = rows_data[1:]
    
    col_widths = []
    n_cols = len(headers)
    for col_idx in range(n_cols):
        max_len = len(headers[col_idx])
        for row in rows:
            if col_idx < len(row):
                max_len = max(max_len, len(row[col_idx]))
        w = max(1.8, min(10.0, max_len * 0.18))
        col_widths.append(w)
        
    printable_width = 15.29
    total_w = sum(col_widths)
    if total_w > printable_width:
        col_widths = [(w / total_w) * printable_width for w in col_widths]
        
    add_professional_table(doc, headers, rows, col_widths=col_widths)

def add_screenshot_placeholder(doc, caption: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    set_para_shading(p, "F3F4F6")
    run = p.add_run(f"[ Screenshot Placeholder: {caption} ]")
    run.font.name = FONT_BODY; run.font.size = Pt(10)
    run.font.italic = True; run.font.color.rgb = MID_GRAY

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    cap_run = cap.add_run(f"Figure: {caption}")
    cap_run.font.name = FONT_BODY; cap_run.font.size = Pt(9)
    cap_run.font.italic = True; cap_run.font.color.rgb = MID_GRAY

# ─── Dynamic Parser ───────────────────────────────────────────────────────────

def parse_markdown_to_docx(doc, md_filepath):
    with open(md_filepath, 'r', encoding='utf-8') as f:
        content = f.read()
        
    start_marker = "# CHAPTER 1: INTRODUCTION"
    idx = content.find(start_marker)
    if idx != -1:
        content = content[idx:]
    
    lines = content.split('\n')
    
    state = 'NORMAL'
    code_lines = []
    table_lines = []
    blockquote_lines = []
    
    def flush_block():
        nonlocal state, code_lines, table_lines, blockquote_lines
        if state == 'CODE_BLOCK' and code_lines:
            add_code_block_parser(doc, code_lines)
            code_lines = []
        elif state == 'TABLE' and table_lines:
            add_table_parser(doc, table_lines)
            table_lines = []
        elif state == 'BLOCKQUOTE' and blockquote_lines:
            add_blockquote_parser(doc, blockquote_lines)
            blockquote_lines = []
        state = 'NORMAL'
        
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        screenshot_match = re.match(r'^\*\*\[SCREENSHOT PLACEHOLDER\s*([^\]]+)\]\*\*$', stripped)
        if screenshot_match:
            flush_block()
            add_screenshot_placeholder(doc, screenshot_match.group(1))
            i += 1
            continue
            
        if state == 'CODE_BLOCK':
            if stripped.startswith('```'):
                flush_block()
            else:
                code_lines.append(line)
        elif state == 'TABLE':
            if stripped.startswith('|'):
                table_lines.append(line)
            else:
                flush_block()
                continue
        elif state == 'BLOCKQUOTE':
            if stripped.startswith('>'):
                blockquote_lines.append(line)
            else:
                flush_block()
                continue
        else: # NORMAL state
            if stripped.startswith('```'):
                state = 'CODE_BLOCK'
                code_lines = []
            elif stripped.startswith('|'):
                state = 'TABLE'
                table_lines = [line]
            elif stripped.startswith('>'):
                state = 'BLOCKQUOTE'
                blockquote_lines = [line]
            elif stripped.startswith('# '):
                styled_heading(doc, stripped[2:].upper(), 1)
            elif stripped.startswith('## '):
                styled_heading(doc, stripped[3:], 2)
            elif stripped.startswith('### '):
                styled_heading(doc, stripped[4:], 3)
            elif stripped.startswith('#### '):
                styled_heading(doc, stripped[5:], 4)
            elif stripped.startswith('---'):
                add_page_break(doc)
            elif stripped.startswith('* ') or stripped.startswith('- ') or stripped.startswith('✓ '):
                indent_spaces = len(line) - len(line.lstrip())
                level = indent_spaces // 2
                prefix = stripped[:2]
                text_content = stripped[2:].strip()
                if prefix == '✓ ':
                    add_bullet_parser(doc, text_content, level, checkmark=True)
                else:
                    add_bullet_parser(doc, text_content, level, checkmark=False)
            elif re.match(r'^\d+\.\s', stripped):
                match = re.match(r'^(\d+)\.\s(.*)', stripped)
                add_numbered_list_parser(doc, match.group(2), int(match.group(1)))
            elif stripped == '':
                pass
            else:
                add_body_para_parser(doc, line)
        i += 1
        
    flush_block()

# ─── Setup Header/Footer ──────────────────────────────────────────────────────

def setup_headers_footers(doc):
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    
    # Running header
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hrun = header_para.add_run("LifeOS: Smart Life Decision Operating System")
    hrun.font.name = FONT_BODY
    hrun.font.size = Pt(8.5)
    hrun.font.color.rgb = MID_GRAY
    
    # Running footer
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    frun = footer_para.add_run("Page ")
    frun.font.name = FONT_BODY
    frun.font.size = Pt(9.5)
    frun.font.color.rgb = MID_GRAY
    add_page_number(frun)

# ─── Main Generator ────────────────────────────────────────────────────────────

def generate():
    md_path = os.path.join(os.path.dirname(__file__), "thesis.md")
    docx_path = os.path.join(os.path.dirname(__file__), "LifeOS_Thesis_Badal_Kumar.docx")
    
    print("Creating LifeOS Thesis Word Document...")
    doc = Document()

    # Global page setup
    section = doc.sections[0]
    set_page_margins(section)
    section.page_width  = Cm(21.0)   # A4
    section.page_height = Cm(29.7)

    configure_styles(doc)
    setup_headers_footers(doc)

    # Front matter
    print("  [1/9] Title page...")
    add_title_page(doc)

    print("  [2/9] Certificate...")
    add_certificate(doc)

    print("  [3/9] Declaration...")
    add_declaration(doc)

    print("  [4/9] Acknowledgement...")
    add_acknowledgement(doc)

    print("  [5/9] Abstract...")
    add_abstract(doc)

    print("  [6/9] Table of Contents...")
    add_toc(doc)
    
    print("  [7/9] List of Figures...")
    add_list_of_figures_page(doc)
    
    print("  [8/9] List of Tables...")
    add_list_of_tables_page(doc)

    # Chapters & Appendices
    print("  [9/9] Parsing chapters and appendices dynamically from thesis.md...")
    parse_markdown_to_docx(doc, md_path)

    # Final closing page
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    r = p.add_run("--- END OF THESIS ---")
    r.font.name = FONT_HEAD; r.font.size = Pt(14); r.font.bold = True; r.font.color.rgb = MID_GRAY

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(16)
    r2 = p2.add_run("Badal Kumar  |  Roll No. 10  |  BCA 6th Semester  |  Session 2023–2026\nDepartment of Computer Application, Patna College, Patna  |  June 2026")
    r2.font.name = FONT_BODY; r2.font.size = Pt(11); r2.font.color.rgb = MID_GRAY

    try:
        doc.save(docx_path)
        size_kb = os.path.getsize(docx_path) // 1024
        print(f"\nDONE: Thesis saved to: {docx_path}")
        print(f"   File size: {size_kb} KB")
    except PermissionError:
        fallback_path = docx_path.replace(".docx", "_Draft.docx")
        doc.save(fallback_path)
        size_kb = os.path.getsize(fallback_path) // 1024
        print(f"\nWARNING: Could not overwrite the original file because it is currently open in another program.")
        print(f"Saved the thesis to: {fallback_path} instead.")
        print(f"   File size: {size_kb} KB")


if __name__ == "__main__":
    generate()
